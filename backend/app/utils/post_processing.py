import re
from docx import Document
from bs4 import BeautifulSoup, NavigableString, Tag
import html


#Formatta una trascrizione grezza di openAI in testo formattato con html
def format_transcription(text: str) -> str:
    """
    Converte una trascrizione grezza in HTML formattato per TipTap.
    Divide il testo in paragrafi rispettando la punteggiatura.
    """
    # Escapa i caratteri HTML
    escaped_text = html.escape(text)

    # Suddivide in paragrafi dopo punti, punti interrogativi, ecc. seguiti da lettera maiuscola
    paragrafi = re.split(r'(?<=[.?!])\s+(?=[A-ZÀ-Ú])', escaped_text)

    # Elimina righe vuote o troppo corte
    paragrafi = [p.strip() for p in paragrafi if len(p.strip()) > 0]

    # Wrappa ogni paragrafo in <p> per TipTap
    html_paragrafi = [f"<p>{p}</p>" for p in paragrafi]

    return "\n".join(html_paragrafi)

#converte una trascrizione con tag html ad una trascrizione ottimizzata per la creazione di un file word
def convert_html_to_word_template(html_text: str) -> Document:
    """
    Converte un testo HTML (es. da TipTap) in un documento Word.
    Preserva la struttura a paragrafi (<p>).
    """
    # Crea un nuovo documento Word
    doc = Document()

    # Parsing dell'HTML
    soup = BeautifulSoup(html_text, "html.parser")

    # Estrae tutti i paragrafi
    paragraphs = soup.find_all("p")

    for p in paragraphs:
        # Decodifica le entità HTML (es: &#x27; -> ')
        clean_text = html.unescape(p.get_text(strip=True))
        doc.add_paragraph(clean_text)

    return doc


def parse_odv_summary(html_text: str) -> dict:
    # 1. Rimuove i tag HTML mantenendo solo il testo puro
    soup = BeautifulSoup(html_text, "html.parser")
    clean_text = soup.get_text(separator="\n")

    # 2. Divide il testo in sezioni usando pattern numerati
    section_pattern = r"(?=^\d\.\s)"  # esempio: "1. "
    sections_raw = re.split(section_pattern, clean_text, flags=re.MULTILINE)
    sections_raw = [s.strip() for s in sections_raw if s.strip()]

    result = {}

    for raw in sections_raw:
        match = re.match(r"^(\d)\.\s(.+)", raw)
        if not match:
            continue

        section_number = int(match.group(1))
        content = raw[match.end():].strip()

        if section_number == 4:
            # Cerca sottosezioni 4.1 Premessa e 4.2 Argomenti trattati
            premessa_match = re.search(
                r"4\.1\s*Premessa\s*(.*?)(?=\n4\.2\s*Argomenti trattati)", content, re.DOTALL | re.IGNORECASE
            )
            argomenti_match = re.search(
                r"4\.2\s*Argomenti trattati\s*(.*)", content, re.DOTALL | re.IGNORECASE
            )

            premessa = premessa_match.group(1).strip() if premessa_match else ""
            argomenti = argomenti_match.group(1).strip() if argomenti_match else ""

            result[4] = {
                "content": content,
                "premessa": premessa,
                "argomenti": argomenti
            }
        else:
            result[section_number] = {
                "content": content
            }

    return result


def fill_odv_template(sections: dict, output_path: str, extra_fields: dict):
    doc = Document("app/templates/template_verbale_odv.docx")

    def replace_all_runs(runs, replacements):
        for run in runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    # Preparazione mappa dei segnaposto
    replacements = {
        "{SEZIONE_1}": sections.get(1, {}).get("content", ""),
        "{SEZIONE_2}": sections.get(2, {}).get("content", ""),
        "{SEZIONE_3}": sections.get(3, {}).get("content", ""),
        "{SEZIONE_4}": sections.get(4, {}).get("content", ""),
        "{SEZIONE_4_PREMESSA}": sections.get(4, {}).get("premessa", ""),
        "{SEZIONE_4_ARGOMENTI}": sections.get(4, {}).get("argomenti", ""),
        "{SEZIONE_5}": sections.get(5, {}).get("content", ""),
        "{SEZIONE_6}": sections.get(6, {}).get("content", ""),
    }

    print(sections.get(4, {}).get("premessa", ""))

    # Aggiunta dei campi extra (es. data, luogo, protocollo...)
    for key, value in extra_fields.items():
        replacements[f"{{{key}}}"] = value

    # Sostituzione nei paragrafi
    for para in doc.paragraphs:
        replace_all_runs(para.runs, replacements)

    # Sostituzione anche nelle celle di eventuali tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_all_runs(para.runs, replacements)

    doc.save(output_path)


def parse_to_tiptap_json(text: str) -> str:
    if "<p>" in text or "<h" in text:
        print("⚠️ Salto parsing: testo già HTML")
        return text  # già in formato HTML valido

    lines = text.split("\n")
    html_parts = []
    heading_pattern = re.compile(r"^\*\*(\d+(\.\d+)?)[\.\)]?\s?(.*?)\*\*$")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        heading_match = heading_pattern.match(line)
        if heading_match:
            level = 2 if "." not in heading_match.group(1) else 3
            title_text = f"{heading_match.group(1)} {heading_match.group(3)}"
            html_parts.append(f"<h{level}>{title_text.strip()}</h{level}>")
            continue

        line = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", line)
        html_parts.append(f"<p>{line}</p>")

    return "\n".join(html_parts)


import re
from bs4 import BeautifulSoup

def _estrai_sezioni_verbale_(html_text: str) -> dict:
    """
    Estrae le sezioni principali da un verbale OdV formattato in HTML.
    Restituisce un dizionario con titoli delle sezioni come chiavi.
    """
    # Rimuove i tag HTML mantenendo solo il testo
    soup = BeautifulSoup(html_text, "html.parser")
    testo_pulito = soup.get_text(separator="\n")

    # Definizione delle intestazioni corrispondenti alle sezioni principali
    pattern_sezioni = {
        "1 Oggetto della riunione": r"1\s+Oggetto della riunione[:]*",
        "2 Processo interessato dal controllo dell’OdV": r"2\s+Processo.*?OdV[:]*",
        "3 Documenti esaminati": r"3\s+Documenti esaminati[:]*",
        "4.1 Premessa": r"4\.1\s+Premessa[:]*",
        "4.2 Argomenti trattati": r"4\.2\s+Argomenti trattati[:]*",
        "5 Considerazioni": r"5\s+Considerazioni[:]*",
        "6 Conclusioni": r"6\s+Conclusioni[:]*"
    }

    # Costruzione regex per suddividere il testo in sezioni
    sezioni_ordinate = list(pattern_sezioni.keys())
    regex = "(" + "|".join(pattern_sezioni.values()) + ")"
    split = re.split(regex, testo_pulito, flags=re.IGNORECASE)

    # Parsing del risultato
    sezioni = {}
    i = 0
    while i < len(split) - 1:
        current = split[i].strip()
        next_block = split[i + 1].strip()
        content = split[i + 2].strip() if i + 2 < len(split) else ""
        for titolo, pattern in pattern_sezioni.items():
            if re.fullmatch(pattern, next_block, flags=re.IGNORECASE):
                sezioni[titolo] = content
                break
        i += 2

    return sezioni


import re
from bs4 import BeautifulSoup

def estrai_sezioni_verbale(html_text: str) -> dict:
    """
    Estrae le sezioni dal verbale in HTML e le restituisce già strutturate
    per la funzione fill_odv_template().
    """
    # Pulisce il testo rimuovendo i tag HTML
    soup = BeautifulSoup(html_text, "html.parser")
    testo = soup.get_text(separator="\n")

    # Sezioni da cercare (titoli nel documento)
    pattern_sezioni = {
        "1 Oggetto della riunione": "1",
        "2 Processo interessato dal controllo dell’OdV": "2",
        "3 Documenti esaminati": "3",
        "4.1 Premessa": "4_premessa",
        "4.2 Argomenti trattati": "4_argomenti",
        "5 Considerazioni": "5",
        "6 Conclusioni": "6"
    }

    # Regex combinata per individuare tutti i titoli di sezione
    regex = "(" + "|".join(re.escape(k) + r":*" for k in pattern_sezioni.keys()) + ")"
    split = re.split(regex, testo, flags=re.IGNORECASE)

    # Parsing del risultato
    sezioni_raw = {}
    i = 0
    while i < len(split) - 1:
        current = split[i].strip()
        next_block = split[i + 1].strip().rstrip(":")
        content = split[i + 2].strip() if i + 2 < len(split) else ""
        sezioni_raw[next_block] = content
        i += 2

    # Mappatura nel formato richiesto da fill_odv_template
    sezioni_mappate = {
        1: {"content": sezioni_raw.get("1 Oggetto della riunione", "")},
        2: {"content": sezioni_raw.get("2 Processo interessato dal controllo dell’OdV", "")},
        3: {"content": sezioni_raw.get("3 Documenti esaminati", "")},
        4: {
            "premessa": sezioni_raw.get("4.1 Premessa", ""),
            "argomenti": sezioni_raw.get("4.2 Argomenti trattati", "")
        },
        5: {"content": sezioni_raw.get("5 Considerazioni", "")},
        6: {"content": sezioni_raw.get("6 Conclusioni", "")}
    }

    return sezioni_mappate
